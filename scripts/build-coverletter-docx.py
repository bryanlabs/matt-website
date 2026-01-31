#!/usr/bin/env python3
"""Convert neuroworks.html cover letter to a styled neuroworks.docx.

Parses the HTML cover letter and builds a Word document that closely
matches the web design: teal header, branded colors, letter formatting.

Usage:
    python scripts/build-coverletter-docx.py
    python scripts/build-coverletter-docx.py neuroworks.html neuroworks.docx
"""

import sys

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# Brand colors from CSS custom properties
PRIMARY = RGBColor(0x0D, 0x6E, 0x6E)
SECONDARY = RGBColor(0x14, 0x91, 0x9B)
ACCENT = RGBColor(0x1A, 0x3A, 0x3A)
DARK = RGBColor(0x2C, 0x3E, 0x50)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY = RGBColor(0x66, 0x66, 0x66)


def set_cell_shading(cell, color_hex):
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def remove_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        f'<w:tblPr {nsdecls("w")}/>'
    )
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0"/>'
        '</w:tblBorders>'
    )
    tbl_pr.append(borders)


def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{int(top * 20)}" w:type="dxa"/>'
        f'  <w:left w:w="{int(left * 20)}" w:type="dxa"/>'
        f'  <w:bottom w:w="{int(bottom * 20)}" w:type="dxa"/>'
        f'  <w:right w:w="{int(right * 20)}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tc_pr.append(margins)


def add_run(paragraph, text, bold=False, italic=False, color=None, size=None):
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    return run


def add_bottom_border(paragraph, color_hex="E8F4F4"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="4" w:color="{color_hex}"/>'
        f'</w:pBdr>'
    )
    p_pr.append(borders)


def parse_cover_letter(html_path):
    with open(html_path, "r") as f:
        soup = BeautifulSoup(f.read(), "lxml")

    data = {}

    header = soup.find("div", class_="header")
    data["name"] = header.find("h1").get_text(strip=True)
    data["title"] = header.find("div", class_="title").get_text(strip=True)

    contact_row = header.find("div", class_="contact-row")
    data["contacts"] = []
    for child in contact_row.children:
        text = child.get_text(strip=True)
        if text:
            data["contacts"].append(text)

    content = soup.find("div", class_="content")
    data["date"] = content.find("div", class_="date").get_text(strip=True)

    recipient = content.find("div", class_="recipient")
    badge = recipient.find("span", class_="company-badge")
    data["badge"] = badge.get_text(strip=True) if badge else ""
    # Get recipient lines (skip badge)
    recipient_lines = []
    for child in recipient.children:
        if child.name == "span" and "company-badge" in (child.get("class") or []):
            continue
        if child.name == "br":
            continue
        text = child.get_text(strip=True) if hasattr(child, "get_text") else str(child).strip()
        if text:
            recipient_lines.append(text)
    data["recipient_lines"] = recipient_lines

    data["salutation"] = content.find("div", class_="salutation").get_text(strip=True)

    body = content.find("div", class_="body")
    data["paragraphs"] = [p.get_text(strip=True) for p in body.find_all("p")]

    closing = content.find("div", class_="closing")
    data["regards"] = closing.find("div", class_="regards").get_text(strip=True)
    data["signature"] = closing.find("div", class_="signature").get_text(strip=True)
    data["credential"] = closing.find("div", class_="credential").get_text(strip=True)

    return data


def build_docx(data, output_path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.6

    # ============================================================
    # HEADER (teal background table)
    # ============================================================
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = 1
    remove_table_borders(header_table)
    cell = header_table.cell(0, 0)
    set_cell_shading(cell, "0D6E6E")
    set_cell_margins(cell, top=14, bottom=14, left=10, right=10)

    tbl = header_table._tbl
    tbl_pr = tbl.tblPr
    for existing in tbl_pr.findall(qn("w:tblW")):
        tbl_pr.remove(existing)
    tbl_pr.append(parse_xml(
        f'<w:tblW {nsdecls("w")} w:w="0" w:type="auto"/>'
    ))

    # Name
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(4)
    add_run(p, data["name"], bold=True, color=WHITE, size=Pt(22))

    # Title
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(0)
    run = add_run(p, data["title"], color=WHITE, size=Pt(11))
    run.font.all_caps = True

    # Contact row
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    for i, contact in enumerate(data["contacts"]):
        if i > 0:
            add_run(p, "    ", color=WHITE, size=Pt(9))
        add_run(p, contact, color=WHITE, size=Pt(9))

    # ============================================================
    # DATE
    # ============================================================
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(14)
    add_run(p, data["date"], color=GRAY, size=Pt(10))

    # ============================================================
    # RECIPIENT
    # ============================================================
    # Company badge
    if data["badge"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = add_run(p, data["badge"], bold=True, color=PRIMARY, size=Pt(9))

    # Recipient lines
    for line in data["recipient_lines"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        # Bold the first line (company name)
        if line == data["recipient_lines"][0]:
            add_run(p, line, bold=True, color=PRIMARY, size=Pt(11))
        else:
            add_run(p, line, color=DARK, size=Pt(11))

    # ============================================================
    # SALUTATION
    # ============================================================
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(10)
    add_run(p, data["salutation"], bold=True, color=ACCENT, size=Pt(11))

    # ============================================================
    # BODY PARAGRAPHS
    # ============================================================
    for para_text in data["paragraphs"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.space_before = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_run(p, para_text, color=DARK, size=Pt(11))

    # ============================================================
    # CLOSING
    # ============================================================
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, data["regards"], italic=True, color=DARK, size=Pt(11))

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_run(p, data["signature"], bold=True, color=PRIMARY, size=Pt(14))

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_run(p, data["credential"], color=GRAY, size=Pt(10))

    # ============================================================
    # FOOTER
    # ============================================================
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_bottom_border(p, "E8F4F4")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    add_run(p, "View Resume: ", color=GRAY, size=Pt(9))
    add_run(p, "mattbfit.net/resume.html", bold=True, color=PRIMARY, size=Pt(9))
    add_run(p, "  |  ", color=GRAY, size=Pt(9))
    add_run(p, "Training Videos: ", color=GRAY, size=Pt(9))
    add_run(p, "mattbfit.net/#showcase", bold=True, color=PRIMARY, size=Pt(9))

    doc.save(output_path)


if __name__ == "__main__":
    html_path = sys.argv[1] if len(sys.argv) > 1 else "neuroworks.html"
    output_path = sys.argv[2] if len(sys.argv) > 2 else "neuroworks.docx"

    data = parse_cover_letter(html_path)
    build_docx(data, output_path)
    print(f"Generated {output_path}")
