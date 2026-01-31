#!/usr/bin/env python3
"""Convert resume.html to a styled resume.docx.

Parses the HTML resume and builds a Word document that closely matches
the web design: teal color scheme, section headers with underlines,
two-column skills grid, formatted job entries, etc.

Usage:
    python scripts/build-resume-docx.py
    python scripts/build-resume-docx.py resume.html resume.docx
"""

import sys
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# Brand colors from CSS custom properties
PRIMARY = RGBColor(0x0D, 0x6E, 0x6E)
SECONDARY = RGBColor(0x14, 0x91, 0x9B)
ACCENT = RGBColor(0x1A, 0x3A, 0x3A)
DARK = RGBColor(0x2C, 0x3E, 0x50)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG = "E8F4F4"


def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def set_paragraph_shading(paragraph, color_hex):
    """Set background shading on a paragraph."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    paragraph._p.get_or_add_pPr().append(shading)


def remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        f'<w:tblPr {nsdecls("w")}/>'
    )
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0"/>'
        '</w:tblBorders>'
    )
    tbl_pr.append(borders)


def set_cell_width(cell, inches):
    """Set a fixed width on a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_width = parse_xml(
        f'<w:tcW {nsdecls("w")} w:w="{int(inches * 1440)}" w:type="dxa"/>'
    )
    tc_pr.append(tc_width)


def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    """Set margins on a table cell (in points)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{int(top * 20)}" w:type="dxa"/>'
        f'  <w:left w:w="{int(left * 20)}" w:type="dxa"/>'
        f'  <w:bottom w:w="{int(bottom * 20)}" w:type="dxa"/>'
        f'  <w:right w:w="{int(right * 20)}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tc_pr.append(margins)


def add_run(paragraph, text, bold=False, color=None, size=None, font_name=None):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = size
    if font_name:
        run.font.name = font_name
    return run


def add_bottom_border(paragraph, color_hex="0D6E6E"):
    """Add a bottom border line under a paragraph."""
    p_pr = paragraph._p.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="4" w:color="{color_hex}"/>'
        f'</w:pBdr>'
    )
    p_pr.append(borders)


def parse_resume(html_path):
    """Parse resume.html and extract structured content."""
    with open(html_path, "r") as f:
        soup = BeautifulSoup(f.read(), "lxml")

    data = {}

    # Header
    header = soup.find("div", class_="header")
    data["name"] = header.find("h1").get_text(strip=True)
    data["title"] = header.find("div", class_="title").get_text(strip=True)

    contact_row = header.find("div", class_="contact-row")
    data["contacts"] = []
    for child in contact_row.children:
        text = child.get_text(strip=True)
        if text:
            data["contacts"].append(text)

    # Skills
    skills_div = soup.find("div", class_="skills-grid")
    data["skills"] = [
        s.get_text(strip=True) for s in skills_div.find_all("div", class_="skill-item")
    ]

    # Jobs
    sections = soup.find_all("div", class_="section")
    data["jobs"] = []
    data["military"] = []
    data["education"] = {}

    for section in sections:
        title_el = section.find("div", class_="section-title")
        if not title_el:
            continue
        title_text = title_el.get_text(strip=True)

        if title_text == "Professional Experience":
            for job_div in section.find_all("div", class_="job"):
                job = _parse_job(job_div)
                data["jobs"].append(job)

        elif title_text == "Military Service":
            for job_div in section.find_all("div", class_="job"):
                job = _parse_job(job_div)
                data["military"].append(job)

        elif title_text == "Education & Certifications":
            two_col = section.find("div", class_="two-column")
            columns = two_col.find_all("div", recursive=False)
            data["education"]["left"] = []
            data["education"]["right"] = []
            for item in columns[0].find_all("div", class_="cert-item"):
                data["education"]["left"].append(item.get_text(" ", strip=True))
            # Also grab the "Education" and "Specialized Training" labels
            for strong in columns[0].find_all("strong"):
                pass  # We'll handle labels inline
            for item in columns[1].find_all("div", class_="cert-item"):
                data["education"]["right"].append(item.get_text(" ", strip=True))

    return data


def _parse_job(job_div):
    """Parse a single job entry."""
    job = {}
    header = job_div.find("div", class_="job-header")

    title_el = header.find("span", class_="job-title")
    company_el = header.find("span", class_="job-company")
    date_el = header.find("span", class_="job-date")

    job["title"] = title_el.get_text(strip=True) if title_el else ""
    job["company"] = company_el.get_text(strip=True) if company_el else ""
    job["date"] = date_el.get_text(strip=True) if date_el else ""

    bullets = job_div.find("ul", class_="job-bullets")
    job["bullets"] = []
    if bullets:
        for li in bullets.find_all("li"):
            job["bullets"].append(li.get_text(strip=True))

    return job


def build_docx(data, output_path):
    """Build a styled Word document from parsed resume data."""
    doc = Document()

    # -- Page margins --
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    page_width_inches = 8.5 - (2 * 2 / 2.54)  # page minus margins in inches

    # -- Default font --
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = DARK
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)

    # ============================================================
    # HEADER BLOCK (using a single-cell table with teal background)
    # ============================================================
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = 1  # center
    remove_table_borders(header_table)
    cell = header_table.cell(0, 0)
    set_cell_shading(cell, "0D6E6E")
    set_cell_margins(cell, top=14, bottom=14, left=10, right=10)

    # Make table span full width
    tbl = header_table._tbl
    tbl_pr = tbl.tblPr
    tbl_width = parse_xml(
        f'<w:tblW {nsdecls("w")} w:w="0" w:type="auto"/>'
    )
    # Remove any existing tblW
    for existing in tbl_pr.findall(qn("w:tblW")):
        tbl_pr.remove(existing)
    tbl_pr.append(tbl_width)

    # Name
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(4)
    add_run(p, data["name"], bold=True, color=WHITE, size=Pt(22))

    # Title
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(0)
    run = add_run(p, data["title"], color=WHITE, size=Pt(11))
    run.font.all_caps = True

    # Contact row
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    for i, contact in enumerate(data["contacts"]):
        if i > 0:
            add_run(p, "    ", color=WHITE, size=Pt(9))
        add_run(p, contact, color=WHITE, size=Pt(9))

    # Spacer after header
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(6)
    spacer.paragraph_format.space_after = Pt(0)

    # ============================================================
    # KEY SKILLS (two-column table)
    # ============================================================
    _add_section_heading(doc, "KEY SKILLS")

    skills = data["skills"]
    rows_needed = (len(skills) + 1) // 2
    skills_table = doc.add_table(rows=rows_needed, cols=2)
    remove_table_borders(skills_table)

    for i, skill in enumerate(skills):
        row_idx = i // 2
        col_idx = i % 2
        cell = skills_table.cell(row_idx, col_idx)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        add_run(p, "\u2713 ", bold=True, color=PRIMARY, size=Pt(10))
        add_run(p, skill, color=DARK, size=Pt(10))

    _add_spacer(doc)

    # ============================================================
    # PROFESSIONAL EXPERIENCE
    # ============================================================
    _add_section_heading(doc, "PROFESSIONAL EXPERIENCE")

    for job in data["jobs"]:
        _add_job_entry(doc, job, page_width_inches)

    # ============================================================
    # MILITARY SERVICE
    # ============================================================
    _add_section_heading(doc, "MILITARY SERVICE")

    for job in data["military"]:
        _add_job_entry(doc, job, page_width_inches)

    # ============================================================
    # EDUCATION & CERTIFICATIONS (two-column)
    # ============================================================
    _add_section_heading(doc, "EDUCATION & CERTIFICATIONS")

    edu_table = doc.add_table(rows=1, cols=2)
    remove_table_borders(edu_table)

    # Left column: Education
    left_cell = edu_table.cell(0, 0)
    p = left_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "Education", bold=True, color=DARK, size=Pt(10))

    for item in data["education"].get("left", []):
        p = left_cell.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.space_before = Pt(1)
        add_run(p, item, color=DARK, size=Pt(9))

    # Right column: Certifications
    right_cell = edu_table.cell(0, 1)
    p = right_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "Certifications", bold=True, color=DARK, size=Pt(10))

    for item in data["education"].get("right", []):
        p = right_cell.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.space_before = Pt(1)
        add_run(p, item, color=DARK, size=Pt(9))

    doc.save(output_path)


def _add_section_heading(doc, text):
    """Add a teal section heading with bottom border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    add_run(p, text, bold=True, color=PRIMARY, size=Pt(12))
    add_bottom_border(p, "0D6E6E")


def _add_spacer(doc, pts=4):
    """Add a small spacer paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pts)


def _add_job_entry(doc, job, page_width_inches):
    """Add a formatted job entry with title | company and date."""
    # Job header line: "Title | Company" on left, date on right
    # Use a table for left/right alignment
    job_table = doc.add_table(rows=1, cols=2)
    remove_table_borders(job_table)

    left_cell = job_table.cell(0, 0)
    right_cell = job_table.cell(0, 1)

    set_cell_width(right_cell, 1.5)

    # Left: Title | Company
    p = left_cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, job["title"], bold=True, color=ACCENT, size=Pt(10))
    if job["company"]:
        add_run(p, " | ", color=DARK, size=Pt(10))
        add_run(p, job["company"], bold=True, color=PRIMARY, size=Pt(10))

    # Right: Date
    p = right_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, job["date"], color=RGBColor(0x66, 0x66, 0x66), size=Pt(9))

    # Bullet points
    for bullet in job["bullets"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.4)
        add_run(p, "\u25b8 ", color=SECONDARY, size=Pt(9))
        add_run(p, bullet, color=DARK, size=Pt(9))


if __name__ == "__main__":
    html_path = sys.argv[1] if len(sys.argv) > 1 else "resume.html"
    output_path = sys.argv[2] if len(sys.argv) > 2 else "resume.docx"

    data = parse_resume(html_path)
    build_docx(data, output_path)
    print(f"Generated {output_path}")
